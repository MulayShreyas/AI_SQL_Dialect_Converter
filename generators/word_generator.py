"""
Word Generator - Creates Word document output files with converted SQL
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Dict
from datetime import datetime
import io


class WordGenerator:
    """
    Generator for creating Word documents with SQL conversion results.
    Uses python-docx for professional document generation.
    """
    
    def __init__(self):
        self.primary_color = RGBColor(102, 126, 234)  # #667eea
        self.success_color = RGBColor(16, 185, 129)    # #10b981
        self.error_color = RGBColor(239, 68, 68)       # #ef4444
        self.gray_color = RGBColor(102, 102, 102)      # #666666
    
    def generate(
        self, 
        results: List[Dict], 
        source_dialect: str, 
        target_dialect: str,
        filename: str = None
    ) -> io.BytesIO:
        """
        Generate a Word document with conversion results.
        
        Args:
            results: List of conversion results
            source_dialect: Source SQL dialect
            target_dialect: Target SQL dialect
            filename: Optional filename (for metadata)
            
        Returns:
            BytesIO buffer containing the Word document
        """
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
        
        # Title
        title = doc.add_heading('SQL Dialect Conversion Report', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f'{source_dialect} → {target_dialect}')
        run.font.size = Pt(14)
        run.font.color.rgb = self.primary_color
        
        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f'Generated on {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        run.font.size = Pt(10)
        run.font.color.rgb = self.gray_color
        
        doc.add_paragraph()
        
        # Summary section
        doc.add_heading('Summary', level=1)
        
        success_count = sum(1 for r in results if r['status'] == 'success')
        error_count = len(results) - success_count
        
        # Summary table
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        summary_data = [
            ('Total Statements', str(len(results))),
            ('Successfully Converted', str(success_count)),
            ('Errors', str(error_count)),
            ('Success Rate', f'{(success_count/len(results)*100):.1f}%' if results else 'N/A')
        ]
        
        for i, (label, value) in enumerate(summary_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # Style the cells
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.size = Pt(11)
        
        doc.add_paragraph()
        
        # Conversion details
        doc.add_heading('Conversion Details', level=1)
        
        for i, result in enumerate(results, 1):
            # Statement header
            header = doc.add_heading(f'Statement {i}', level=2)
            
            # Status indicator
            status_para = doc.add_paragraph()
            if result['status'] == 'success':
                run = status_para.add_run('✓ Successfully Converted')
                run.font.color.rgb = self.success_color
            else:
                run = status_para.add_run('✗ Conversion Failed')
                run.font.color.rgb = self.error_color
            run.font.bold = True
            
            # Original SQL
            doc.add_heading('Original SQL:', level=3)
            self._add_code_block(doc, result['original'])
            
            # Converted SQL or Error
            if result['status'] == 'success':
                doc.add_heading('Converted SQL:', level=3)
                self._add_code_block(doc, result['converted'])
                
                # Notes
                if result.get('notes'):
                    notes_para = doc.add_paragraph()
                    run = notes_para.add_run('Notes: ')
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run = notes_para.add_run(result['notes'])
                    run.font.size = Pt(10)
                    run.font.italic = True
                    run.font.color.rgb = self.gray_color
            else:
                doc.add_heading('Error:', level=3)
                error_para = doc.add_paragraph()
                run = error_para.add_run(result.get('notes', 'Unknown error'))
                run.font.color.rgb = self.error_color
            
            # Add spacing between statements
            doc.add_paragraph()
            
            # Add page break every 2 statements
            if i % 2 == 0 and i < len(results):
                doc.add_page_break()
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    def _add_code_block(self, doc: Document, code: str):
        """Add a formatted code block to the document."""
        if not code:
            code = "(empty)"
        
        # Create a table with single cell for code block effect
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        
        cell = table.rows[0].cells[0]
        cell.text = code
        
        # Style the code
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
        
        # Set cell shading (light gray background)
        self._set_cell_shading(cell, 'F5F5F5')
    
    def _set_cell_shading(self, cell, color: str):
        """Set the background color of a cell."""
        cell_properties = cell._tc.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        cell_properties.append(shading)
